# -*- coding: utf-8 -*-
"""
Beta-Product Representational Network Analysis


Outputs:
  - Time-course plots per model
  - Glass brain SVGs per time window per model (right sagittal view)
  - APA-style Word tables (one per model) listing all FDR-significant
    edges

Author: Ioannis Ntoumanis
Created: 2026
"""

import os
import numpy as np
import nibabel as nib
import h5py
import pickle
import time
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.colors as mpl_colors
from scipy.spatial.distance import squareform
from scipy.stats import rankdata
from nilearn import plotting as npl
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# ===========================================================================
# CONFIGURATION
# ===========================================================================

SOURCE_AVG_FILE = r"Source data/source_avg.npy"
VALID_MASK_FILE = r"Source data/valid_mask.npy"
TIME_WIN_FILE   = r"Source data/time_windows.npy"
MNI_LOC_FILE    = r"Source data/grid_loc_mni.mat"
MODEL_RDM_DIR   = r"Results/Model RDMs"
AAL_ATLAS_FILE  = r"Code/AAL3v1.nii"
OUTPUT_DIR      = r"Results/Connectivity"

MODEL_NAMES  = ["rdm_valence", "rdm_arousal"]
MODEL_LABELS = ["Valence", "Arousal"]

N_PERM      = 10000
RANDOM_SEED = 42
FDR_Q       = 0.05

# Fixed colormap range for glass brain edges
EDGE_VMIN = 0
EDGE_VMAX =  0.07

ROI_KEYWORDS = {
    'Angular':     ['Angular_L',          'Angular_R'],
    'MTG':         ['Temporal_Mid_L',     'Temporal_Mid_R'],
    'STG':         ['Temporal_Sup_L',     'Temporal_Sup_R'],
    'IOG':         ['Occipital_Inf_L',    'Occipital_Inf_R'],
    'Lingual':     ['Lingual_L',          'Lingual_R'],
    'Fusiform':    ['Fusiform_L',         'Fusiform_R'],
    'PCC':         ['Cingulate_Post_L',   'Cingulate_Post_R'],
    'Precuneus':   ['Precuneus_L',        'Precuneus_R'],
    'MCC':         ['Cingulate_Mid_L',    'Cingulate_Mid_R'],
    'ACC':         ['ACC_sup_L',          'ACC_sup_R'],
    'Amygdala':    ['Amygdala_L',         'Amygdala_R'],
    'Hippocampus': ['Hippocampus_L',      'Hippocampus_R'],
    'dlPFC':       ['Frontal_Sup_2_L',    'Frontal_Sup_2_R'],
    'MFG':         ['Frontal_Mid_2_L',    'Frontal_Mid_2_R'],
    'IFG':         ['Frontal_Inf_Oper_L', 'Frontal_Inf_Oper_R'],
    'OFC':         ['Frontal_Med_Orb_L',  'Frontal_Med_Orb_R'],
    'Insula':      ['Insula_L',           'Insula_R'],
}

# ===========================================================================
# HELPERS
# ===========================================================================

def rank_vector(v):
    return rankdata(v).astype(np.float64)

def geom_mean_conn(b1, b2):
    """sign(b1*b2) * sqrt(|b1*b2|) — geometric mean preserving sign."""
    product = b1 * b2
    return np.sign(product) * np.sqrt(np.abs(product))

def fdr_bh(p_values_flat, q=0.05):
    n     = len(p_values_flat)
    valid = ~np.isnan(p_values_flat)
    p_v   = p_values_flat[valid]
    nv    = len(p_v)
    if nv == 0:
        return np.zeros(n, dtype=bool), 0.0
    sidx   = np.argsort(p_v)
    sp     = p_v[sidx]
    crit   = np.arange(1, nv + 1) / nv * q
    below  = sp <= crit
    if not np.any(below):
        return np.zeros(n, dtype=bool), 0.0
    thr    = sp[np.max(np.where(below)[0])]
    sig    = np.zeros(n, dtype=bool)
    sig[valid] = p_values_flat[valid] <= thr
    return sig, thr

def assign_grid_to_aal(grid_mm, aal_img, aal_labels, aal_indices, roi_dict,
                        radius_mm=10.0):
    aal_data   = aal_img.get_fdata()
    aal_affine = aal_img.affine
    inv_affine = np.linalg.inv(aal_affine)

    index_to_roi = {}
    for roi_name, label_substrings in roi_dict.items():
        for substr in label_substrings:
            for i, label in enumerate(aal_labels):
                if label == substr or substr in label:
                    index_to_roi[int(aal_indices[i])] = roi_name

    print(f"  AAL indices mapped to ROIs: {len(index_to_roi)}")
    roi_assignments = {name: [] for name in roi_dict.keys()}

    for g in range(len(grid_mm)):
        mni_coord = np.array([grid_mm[g, 0], grid_mm[g, 1], grid_mm[g, 2], 1.0])
        vox       = inv_affine @ mni_coord
        vi, vj, vk = int(round(vox[0])), int(round(vox[1])), int(round(vox[2]))

        assigned = False
        if (0 <= vi < aal_data.shape[0] and
            0 <= vj < aal_data.shape[1] and
            0 <= vk < aal_data.shape[2]):
            aal_idx = int(aal_data[vi, vj, vk])
            if aal_idx in index_to_roi:
                roi_assignments[index_to_roi[aal_idx]].append(g)
                assigned = True

        if not assigned:
            search_range = int(np.ceil(radius_mm / np.abs(aal_affine[0, 0])))
            found = False
            for di in range(-search_range, search_range + 1):
                if found: break
                for dj in range(-search_range, search_range + 1):
                    if found: break
                    for dk in range(-search_range, search_range + 1):
                        ni, nj, nk = vi + di, vj + dj, vk + dk
                        if (0 <= ni < aal_data.shape[0] and
                            0 <= nj < aal_data.shape[1] and
                            0 <= nk < aal_data.shape[2]):
                            aal_idx = int(aal_data[ni, nj, nk])
                            if aal_idx in index_to_roi:
                                vox_mm = aal_affine @ np.array([ni, nj, nk, 1.0])
                                dist   = np.sqrt(np.sum((grid_mm[g] - vox_mm[:3])**2))
                                if dist <= radius_mm:
                                    roi_assignments[index_to_roi[aal_idx]].append(g)
                                    found = True
                                    break

    for name in roi_assignments:
        roi_assignments[name] = np.array(roi_assignments[name], dtype=int)
        print(f"    {name}: {len(roi_assignments[name])} grid points")

    return roi_assignments

def make_sig_adjacency(m_idx, t, obs_conn, fdr_sig, n_rois):
    mat = obs_conn[m_idx, :, :, t].copy()
    sig = fdr_sig[m_idx, :, :, t]
    mat[~sig] = 0.0
    np.fill_diagonal(mat, 0.0)
    return mat

# ===========================================================================
# MAIN
# ===========================================================================
t_start = time.time()
os.makedirs(OUTPUT_DIR, exist_ok=True)

print("=" * 60)
print("BETA-PRODUCT REPRESENTATIONAL NETWORK")
print(f"  ROIs: {len(ROI_KEYWORDS)},  Permutations: {N_PERM}")
print("=" * 60)

# ── 1. Load data ──────────────────────────────────────────────────────────────
print("\nStep 1: Loading data")
source_avg   = np.load(SOURCE_AVG_FILE)
valid_mask   = np.load(VALID_MASK_FILE)
time_windows = np.load(TIME_WIN_FILE)

n_stimuli, n_grid, n_time = source_avg.shape
valid_idx = np.where(valid_mask)[0]
n_valid   = len(valid_idx)
n_pairs   = n_valid * (n_valid - 1) // 2

print(f"  Source: {source_avg.shape},  Valid stimuli: {n_valid}")

with h5py.File(MNI_LOC_FILE, 'r') as f:
    grid_mni = np.array(f['grid_mni']).T
grid_mm = grid_mni * 1000

n_models = len(MODEL_NAMES)
model_rdms_trimmed = []
for name in MODEL_NAMES:
    rdm      = np.load(os.path.join(MODEL_RDM_DIR, f"{name}.npy"))
    rdm_trim = rdm[np.ix_(valid_idx, valid_idx)]
    model_rdms_trimmed.append(rdm_trim)

# ── 2. Design matrices ────────────────────────────────────────────────────────
print("\nStep 2: Building design matrices")
obs_model_vecs = []
for m_idx in range(n_models):
    vec = squareform(model_rdms_trimmed[m_idx], checks=False)
    rv  = rank_vector(vec);  rv -= rv.mean()
    obs_model_vecs.append(rv)

X_obs       = np.column_stack(obs_model_vecs + [np.ones(n_pairs)])
Xt_obs      = X_obs.T
XtX_inv_obs = np.linalg.inv(Xt_obs @ X_obs)

print(f"  Design matrix: {X_obs.shape}")
print(f"  Building {N_PERM} permuted design matrices ...", end=" ", flush=True)

rng = np.random.default_rng(RANDOM_SEED)
perm_orders = np.zeros((N_PERM, n_valid), dtype=np.int64)
for p in range(N_PERM):
    perm_orders[p] = rng.permutation(n_valid)

Xt_stack      = np.zeros((N_PERM, n_models + 1, n_pairs), dtype=np.float64)
XtX_inv_stack = np.zeros((N_PERM, n_models + 1, n_models + 1), dtype=np.float64)

for p in range(N_PERM):
    cols = []
    for m_idx in range(n_models):
        rdm_perm = model_rdms_trimmed[m_idx][np.ix_(perm_orders[p], perm_orders[p])]
        vec = rank_vector(squareform(rdm_perm, checks=False));  vec -= vec.mean()
        cols.append(vec)
    cols.append(np.ones(n_pairs))
    X_perm            = np.column_stack(cols)
    Xt_stack[p]       = X_perm.T
    XtX_inv_stack[p]  = np.linalg.inv(X_perm.T @ X_perm)
print("done")

# ── 3. ROI assignment ─────────────────────────────────────────────────────────
print("\nStep 3: Assigning grid points to AAL3 ROIs")
aal_img = nib.load(AAL_ATLAS_FILE)
aal_dir = os.path.dirname(AAL_ATLAS_FILE)
labels = [];  indices = []
for label_file in ['AAL3v1.nii.txt', 'AAL3v1_1mm.nii.txt',
                   'ROI_MNI_V7.txt', 'AAL3v1.xml', 'AAL3_labels.txt']:
    lf = os.path.join(aal_dir, label_file)
    if os.path.exists(lf):
        print(f"  Found label file: {lf}")
        with open(lf, 'r') as f:
            for line in f:
                parts = line.strip().split()
                if len(parts) >= 2:
                    try:
                        indices.append(int(parts[0]))
                        labels.append(parts[1])
                    except ValueError:
                        continue
        break

roi_assignments = assign_grid_to_aal(grid_mm, aal_img, labels, indices,
                                      ROI_KEYWORDS, radius_mm=10.0)
active_rois = {n: idx for n, idx in roi_assignments.items() if len(idx) >= 5}
roi_names   = [r for r in ROI_KEYWORDS.keys() if r in active_rois]
n_rois      = len(roi_names)
print(f"\n  Active ROIs: {n_rois}")

# ── 4. Partial RSA betas for all ROIs × time windows ─────────────────────────
print("\nStep 4: Computing partial RSA betas")
source_valid = source_avg[valid_idx, :, :]

obs_betas  = np.full((n_rois, n_time, n_models), np.nan)
perm_betas = np.full((n_rois, n_time, N_PERM, n_models), np.nan, dtype=np.float32)

for r_idx, roi_name in enumerate(roi_names):
    grid_indices = active_rois[roi_name]
    print(f"  {roi_name} ({len(grid_indices)} grid pts) ...", end=" ", flush=True)

    for t in range(n_time):
        patterns    = source_valid[:, grid_indices, t]
        patterns_zm = patterns - patterns.mean(axis=1, keepdims=True)
        norms       = np.sqrt(np.sum(patterns_zm ** 2, axis=1))
        if np.any(norms < 1e-12):
            continue
        patterns_n  = patterns_zm / norms[:, np.newaxis]
        corr_mat    = patterns_n @ patterns_n.T
        np.clip(corr_mat, -1.0, 1.0, out=corr_mat)
        neural_vec  = squareform(1.0 - corr_mat, checks=False)
        nr          = rank_vector(neural_vec);  nr -= nr.mean()

        b = XtX_inv_obs @ (Xt_obs @ nr)
        obs_betas[r_idx, t, :] = b[:n_models]

        XtY_all        = np.einsum('ijk,k->ij', Xt_stack, nr)
        perm_betas_all = np.einsum('ijk,ik->ij', XtX_inv_stack, XtY_all)
        perm_betas[r_idx, t, :, :] = perm_betas_all[:, :n_models].astype(np.float32)

    print("done")

# ── 5. Co-representational strength ───────────────────────────────────────────
print("\nStep 5: Computing Co-representational strength")
triu_r1, triu_r2 = np.triu_indices(n_rois, k=1)
n_roi_pairs      = len(triu_r1)

obs_conn  = np.full((n_models, n_rois, n_rois, n_time), np.nan)
perm_conn = np.full((n_models, n_roi_pairs, n_time, N_PERM), np.nan, dtype=np.float32)

for m_idx in range(n_models):
    for t in range(n_time):
        b  = obs_betas[:, t, m_idx]
        pb = perm_betas[:, t, :, m_idx]

        for pair_idx, (r1, r2) in enumerate(zip(triu_r1, triu_r2)):
            gm = geom_mean_conn(b[r1], b[r2])
            obs_conn[m_idx, r1, r2, t] = gm
            obs_conn[m_idx, r2, r1, t] = gm

            perm_gm = geom_mean_conn(pb[r1, :], pb[r2, :])
            perm_conn[m_idx, pair_idx, t, :] = perm_gm

print("  Done")

# ── 6. P-values and FDR per model ─────────────────────────────────────────────
print("\nStep 6: P-values and FDR correction (separately per model)")
p_values     = np.full((n_models, n_rois, n_rois, n_time), np.nan)
p_fdr        = np.full((n_models, n_rois, n_rois, n_time), np.nan)
fdr_sig      = np.zeros((n_models, n_rois, n_rois, n_time), dtype=bool)
fdr_threshold_per_model = np.full(n_models, np.nan)

for m_idx in range(n_models):
    p_flat       = []
    flat_indices = []

    for pair_idx, (r1, r2) in enumerate(zip(triu_r1, triu_r2)):
        for t in range(n_time):
            obs_val   = obs_conn[m_idx, r1, r2, t]
            null_dist = perm_conn[m_idx, pair_idx, t, :]
            if np.isnan(obs_val) or np.all(np.isnan(null_dist)):
                p = np.nan
            else:
                # Two-tailed: |null| >= |observed|
                p = np.mean(np.abs(null_dist) >= np.abs(obs_val))
            p_values[m_idx, r1, r2, t] = p
            p_values[m_idx, r2, r1, t] = p
            p_flat.append(p)
            flat_indices.append((r1, r2, t))

    p_flat   = np.array(p_flat)
    sig_flat, threshold = fdr_bh(p_flat, q=FDR_Q)
    fdr_threshold_per_model[m_idx] = threshold

    # Compute FDR-adjusted p (BH step-up) for reporting
    valid       = ~np.isnan(p_flat)
    p_v         = p_flat[valid]
    nv          = len(p_v)
    order       = np.argsort(p_v)
    ranked_p    = p_v[order]
    adj_ranked  = ranked_p * nv / (np.arange(1, nv + 1))
    # Enforce monotonicity (BH adjusted p is min of all adj_p at >= rank)
    adj_ranked  = np.minimum.accumulate(adj_ranked[::-1])[::-1]
    adj_ranked  = np.clip(adj_ranked, 0, 1)
    p_adj_flat  = np.full_like(p_flat, np.nan)
    inv_order   = np.empty_like(order)
    inv_order[order] = np.arange(nv)
    p_adj_flat[valid] = adj_ranked[inv_order]

    for k, (r1, r2, t) in enumerate(flat_indices):
        fdr_sig[m_idx, r1, r2, t] = sig_flat[k]
        fdr_sig[m_idx, r2, r1, t] = sig_flat[k]
        p_fdr[m_idx, r1, r2, t]   = p_adj_flat[k]
        p_fdr[m_idx, r2, r1, t]   = p_adj_flat[k]

    print(f"  {MODEL_LABELS[m_idx]}: {sig_flat.sum()} / {len(p_flat)} "
          f"significant (FDR q<{FDR_Q}, p threshold={threshold:.4f})")

for m_idx in range(n_models):
    for t in range(n_time):
        b = obs_betas[:, t, m_idx]
        # Build a mask: True where both betas are in the same direction
        pos_mask = ((b[:, None] > 0) & (b[None, :] > 0)) | ((b[:, None] < 0) & (b[None, :] < 0))
        # Apply to fdr_sig
        fdr_sig[m_idx, :, :, t] &= pos_mask

    n_kept = fdr_sig[m_idx].sum() // 2  # upper-triangle count
    print(f"  {MODEL_LABELS[m_idx]}: {n_kept} positive-positive edges retained")
    
# ── 7. Save arrays ────────────────────────────────────────────────────────────
print(f"\nStep 7: Saving results to {OUTPUT_DIR}/")
np.save(os.path.join(OUTPUT_DIR, "obs_betas.npy"),    obs_betas)
np.save(os.path.join(OUTPUT_DIR, "obs_conn.npy"),     obs_conn)
np.save(os.path.join(OUTPUT_DIR, "p_values.npy"),     p_values)
np.save(os.path.join(OUTPUT_DIR, "p_fdr.npy"),        p_fdr)
np.save(os.path.join(OUTPUT_DIR, "fdr_sig.npy"),      fdr_sig)
np.save(os.path.join(OUTPUT_DIR, "time_windows.npy"), time_windows)

with open(os.path.join(OUTPUT_DIR, "info.pkl"), 'wb') as f:
    pickle.dump({'roi_names':    roi_names,
                 'model_names':  MODEL_NAMES,
                 'model_labels': MODEL_LABELS,
                 'n_perm':       N_PERM,
                 'fdr_q':        FDR_Q,
                 'random_seed':  RANDOM_SEED}, f)
print("  Saved.")

# ===========================================================================
# 8. GLASS BRAIN NETWORK FIGURES
# ===========================================================================
print("\nStep 9: Glass brain figures")

# ── Compute centroids from AAL atlas (right hemisphere only) ──────────────────
ROI_VIZ_LABEL = {
    'dlPFC': 'Frontal_Sup_2_R',
}

label_to_idx = {label: int(indices[i]) for i, label in enumerate(labels)}

roi_centroids = np.zeros((n_rois, 3))
roi_n_voxels  = np.zeros(n_rois)

for r_idx, roi_name in enumerate(roi_names):
    if roi_name in ROI_VIZ_LABEL:
        right_label = ROI_VIZ_LABEL[roi_name]
    else:
        right_label = ROI_KEYWORDS[roi_name][1]

    if right_label not in label_to_idx:
        print(f"  WARNING: {right_label} not in atlas labels")
        continue
    aal_code  = label_to_idx[right_label]
    voxel_idx = np.argwhere(aal_img.get_fdata() == aal_code)
    if len(voxel_idx) == 0:
        continue
    voxel_h    = np.column_stack([voxel_idx, np.ones(len(voxel_idx))])
    mni_coords = (aal_img.affine @ voxel_h.T).T[:, :3]
    roi_centroids[r_idx] = mni_coords.mean(axis=0)
    roi_n_voxels[r_idx]  = len(voxel_idx)

cmap_gb = {'Valence': plt.get_cmap('Reds'), 'Arousal': plt.get_cmap('Blues')}
norm_gb = mpl_colors.TwoSlopeNorm(vmin=EDGE_VMIN, vcenter=(EDGE_VMAX+EDGE_VMIN)/2, vmax=EDGE_VMAX)

def draw_glass_brain(adj_mat, model_label, tw_s, tw_e):
    n_sig_conn = int(np.sum(adj_mat != 0) // 2)

    involved = np.any(adj_mat != 0, axis=0)
    if involved.sum() == 0:
        fig, ax = plt.subplots(figsize=(5, 4))
        ax.text(0.5, 0.5,
                f"{model_label}  |  {tw_s}–{tw_e} ms  |  0 edges",
                ha='center', va='center', fontsize=10, fontweight='bold',
                transform=ax.transAxes)
        ax.axis('off')
        fname = f"glassbrain_{model_label.lower()}_{tw_s}-{tw_e}ms.svg"
        fig.savefig(os.path.join(OUTPUT_DIR, fname), format='svg', bbox_inches='tight')
        plt.close(fig)
        return fname

    coords_used = roi_centroids[involved]
    adj_used    = adj_mat[np.ix_(involved, involved)]

    display = npl.plot_connectome(
        adj_used,
        coords_used,
        node_color='#2c3e50',
        node_size=100,
        display_mode='r',
        alpha=0.6,
        edge_vmin=EDGE_VMIN,
        edge_vmax=EDGE_VMAX,
        edge_cmap=cmap_gb[model_label],#plt.get_cmap('Reds'),
        edge_kwargs={'linewidth': 2.5},
        colorbar=True,
        annotate=False,
    )
    display.title(
        text=f"{model_label}  |  {tw_s}–{tw_e} ms  |  {n_sig_conn} edges",
        color='black', bgcolor='white', size=10, fontweight='bold'
    )

    fname  = f"glassbrain_{model_label.lower()}_{tw_s}-{tw_e}ms.svg"
    display.savefig(os.path.join(OUTPUT_DIR, fname))
    display.close()
    return fname

for m_idx, model_label in enumerate(MODEL_LABELS):
    print(f"\n  {model_label}:")
    for t in range(n_time):
        tw_s = int(round(time_windows[t, 0] * 1000))
        tw_e = int(round(time_windows[t, 1] * 1000))
        adj  = make_sig_adjacency(m_idx, t, obs_conn, fdr_sig, n_rois)
        n_sig = np.sum(adj != 0) // 2
        fname = draw_glass_brain(adj, model_label, tw_s, tw_e)
        print(f"    {tw_s}–{tw_e} ms: {n_sig} edges → {fname}")

# ===========================================================================
# 9. APA-STYLE WORD TABLES
# ===========================================================================
print("\nStep 9: APA-style Word tables")

def format_p_apa(p):
    """APA style p-value: '< .001' for tiny p, otherwise '.XXX' without leading 0."""
    if np.isnan(p):
        return "—"
    if p < 0.001:
        return "< .001"
    s = f"{p:.3f}"
    if s.startswith("0."):
        s = s[1:]
    return s

def format_gm_apa(gm):
    """APA style geometric mean: 2-3 sig figs, no leading zero."""
    if np.isnan(gm):
        return "—"
    s = f"{gm:.3f}"
    if s.startswith("0."):
        s = s[1:]
    elif s.startswith("-0."):
        s = "-" + s[2:]
    return s

def set_cell_text(cell, text, bold=False, italic=False, font_size=11,
                   align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(font_size)
    run.bold      = bold
    run.italic    = italic
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def build_table_doc(model_label, m_idx):
    """Build APA-style table listing all FDR-significant edges."""
    # Gather rows
    rows = []
    for t in range(n_time):
        for pair_idx, (r1, r2) in enumerate(zip(triu_r1, triu_r2)):
            if fdr_sig[m_idx, r1, r2, t]:
                tw_s = int(round(time_windows[t, 0] * 1000))
                tw_e = int(round(time_windows[t, 1] * 1000))
                rows.append({
                    'time_label': f"{tw_s}–{tw_e}",
                    't_start':    tw_s,
                    'roi_pair':   f"{roi_names[r1]} – {roi_names[r2]}",
                    'gm':         obs_conn[m_idx, r1, r2, t],
                    'p_raw':      p_values[m_idx, r1, r2, t],
                })

    # Sort by time then ROI pair name
    rows.sort(key=lambda r: (r['t_start'], r['roi_pair']))

    # ── Build document ─────────────────────────────────────────────────────
    doc = Document()

    # Page margins (APA: 1 inch all around)
    section = doc.sections[0]
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

    # Table number (APA: bold, plain font)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Table {m_idx + 1}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = True

    # Title (APA: italic)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(
        f"FDR-significant edges of the {model_label.lower()} representational network"
    )
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = True

    # Create table
    headers = ['Time window (ms)', 'ROI pair',
               'Geometric mean', 'p']
    n_cols  = len(headers)
    n_data_rows = len(rows)
    table = doc.add_table(rows=n_data_rows + 1, cols=n_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.autofit = True

    # Header row
    for c, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[c], h,
                       bold=False, italic=True,
                       font_size=11, align='center')

    # Data rows
    for ri, r in enumerate(rows, start=1):
        set_cell_text(table.rows[ri].cells[0], r['time_label'],
                       align='center', font_size=11)
        set_cell_text(table.rows[ri].cells[1], r['roi_pair'],
                       align='left', font_size=11)
        set_cell_text(table.rows[ri].cells[2], format_gm_apa(r['gm']),
                       align='center', font_size=11)
        set_cell_text(table.rows[ri].cells[3], format_p_apa(r['p_raw']),
                       align='center', font_size=11)

    # APA-style borders: only top, header bottom, and bottom of table
    # Remove all borders first, then add only the three horizontal ones
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_cell_border(cell, **kwargs):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        for edge, attrs in kwargs.items():
            tag = qn(f'w:{edge}')
            elem = tcBorders.find(tag)
            if elem is None:
                elem = OxmlElement(f'w:{edge}')
                tcBorders.append(elem)
            for k, v in attrs.items():
                elem.set(qn(f'w:{k}'), v)

    none_border = {'val': 'nil'}
    single      = {'val': 'single', 'sz': '8', 'color': '000000'}

    for ri in range(len(table.rows)):
        for c in range(n_cols):
            cell = table.rows[ri].cells[c]
            # Default: no borders
            set_cell_border(cell,
                            top=none_border, bottom=none_border,
                            left=none_border, right=none_border)
            # Top border on first row
            if ri == 0:
                set_cell_border(cell, top=single)
            # Bottom border under header (separator)
            if ri == 0:
                set_cell_border(cell, bottom=single)
            # Bottom border on last row
            if ri == len(table.rows) - 1:
                set_cell_border(cell, bottom=single)

    # Note below table (APA: italic "Note." then plain text)
    doc.add_paragraph()
    note = doc.add_paragraph()
    note_run_italic = note.add_run("Note. ")
    note_run_italic.font.name = 'Times New Roman'
    note_run_italic.font.size = Pt(10)
    note_run_italic.italic = True
    note_run = note.add_run(
    f"Corepresentation strength S = sign(β_i × β_j) × √(β_i × β_j), where β_i and β_j are "
    f"partial RSA beta coefficients for the {model_label.lower()} model at "
    f"the same time window for the two ROIs. Only edges where both regions "
    f"showed positive (canonical) encoding (β > 0) are reported, to avoid "
    f"suppression artifacts arising from the valence–arousal correlation in "
    f"the stimulus set. Significance assessed with {N_PERM:,} stimulus-label "
    f"permutations; raw permutation p values are reported, and only edges "
    f"surviving FDR correction (Benjamini–Hochberg, q < .{int(FDR_Q*100):02d}) "
    f"across all ROI pairs × time windows separately per model are listed."
    )
    note_run.font.name = 'Times New Roman'
    note_run.font.size = Pt(10)

    out_doc = os.path.join(OUTPUT_DIR,
                           f"table_significant_{model_label.lower()}.docx")
    doc.save(out_doc)
    return out_doc, n_data_rows

for m_idx, model_label in enumerate(MODEL_LABELS):
    out_doc, n_rows = build_table_doc(model_label, m_idx)
    print(f"  {model_label}: {n_rows} significant edges → {out_doc}")

total_time = time.time() - t_start
print(f"\n{'='*60}")
print(f"DONE. Runtime: {total_time/60:.1f} min")
print(f"{'='*60}")